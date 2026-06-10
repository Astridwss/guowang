# 文件路径：services/document_splitter.py
import os
import sys
import io
import re
import pandas as pd
from docx import Document
from PIL import Image

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
# 跨层调用底层公共工具，服务自身保持极简
from utils.text_processor import TextProcessor

class DocumentSplitterService:
    """
    文档提取与切分综合服务 (ETL - Extract)
    职责：统一处理底层源文件，支持读取 Excel 结构化表格 和 Word 多模态图文手册。
    """
    def __init__(self, output_img_dir:str = None):
        # Word 图像输出目录
        self.output_img_dir = output_img_dir
        print("[DocumentSplitterService] 初始化文档统一解析与切分组件...")

    # ============================================================
    # 1. Excel 结构化表格处理逻辑
    # ============================================================
    def process_faq_excel(self, excel_file_path: str) -> tuple:
        """
        处理存量 FAQ Excel 问答对表格
        读取--问题清洗--返回[问题，答案，业务域]
        返回: (valid_texts文本列表, payloads元数据列表)
        """
        if not os.path.exists(excel_file_path):
            print(f"[DocumentSplitterService] ❌ 找不到 Excel 文件: {excel_file_path}")
            return [], []

        print(f"[DocumentSplitterService] 📦 正在加载 Excel: {excel_file_path}")
        df = pd.read_excel(excel_file_path)
        # 去重逻辑
        df.drop_duplicates(subset=["问题"], keep='first', inplace=True)
        
        valid_texts = []
        payloads = []

        for _, row in df.iterrows():
            # 核心：调用公共清洗规则，确保与检索端对齐
            clean_q = TextProcessor.clean_text(str(row.get('问题', '')))
            if not clean_q:
                continue
                
            valid_texts.append(clean_q)
            payloads.append({
                "question": clean_q,
                "answer": str(row.get('答案', '')),
                "domain": str(row.get('所属系统', '未知'))
            })
            
        print(f"[DocumentSplitterService]  Excel 解析完毕，提取 {len(valid_texts)} 条有效数据。")
        return valid_texts, payloads


    # ============================================================
    # 2. Word 多模态图文手册物理切分逻辑
    # ============================================================
    def process_word_docx(self, docx_path: str) -> list:
        """
        处理非结构化 Word 图文手册
        返回: 按标题切分好的多模态文档块列表 (List of Lists)
        """
        if not os.path.exists(docx_path):
            print(f"[DocumentSplitterService] ❌ 找不到 Word 文档: {docx_path}")
            return []
            
        print(f"[DocumentSplitterService] 📄 正在解析 Word 文档: {os.path.basename(docx_path)}")
        raw_elements = self._extract_word_elements(docx_path)
        chunks = self._split_by_heading(raw_elements)
        
        print(f"[DocumentSplitterService]  Word 解析完毕，切分为 {len(chunks)} 个逻辑块。")
        return chunks

    def _extract_word_elements(self, docx_path: str) -> list:
        """底层方法：遍历 Word 提取文本、表格与内嵌图片"""
        doc = Document(docx_path)
        elements = []
        rels = doc.part._rels
        img_id = 0

        def get_heading_level(style_name):
            match = re.match(r'Heading (\d+)', style_name)
            return f"h{match.group(1)}" if match else None

        for block in doc.element.body:
            if block.tag.endswith('}p'): 
                para = next(p for p in doc.paragraphs if p._element == block)
                if para.text.strip():
                    elements.append({
                        "type": "text", 
                        "text": para.text.strip(), 
                        "heading": get_heading_level(para.style.name)
                    })
                
                # 提取并保存内嵌图片
                for run in para.runs:
                    match = re.search(r'r:embed="(rId\d+)"', run._element.xml)
                    if match:
                        rId = match.group(1)
                        if rId in rels:
                            image_part = rels[rId].target_part
                            image = Image.open(io.BytesIO(image_part.blob))
                            
                            doc_name = os.path.splitext(os.path.basename(docx_path))[0]
                            img_filename = f"{doc_name}_image_{img_id}.png"
                            img_path = os.path.join(self.output_img_dir, img_filename)
                            
                            image.save(img_path)
                            elements.append({
                                "type": "image", 
                                "image": img_filename, 
                                "heading": None
                            })
                            img_id += 1

            elif block.tag.endswith('}tbl'): 
                tbl = next(t for t in doc.tables if t._element == block)
                table_text = []
                for row in tbl.rows:
                    row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    table_text.append('\t'.join(row_text))
                elements.append({
                    "type": "text", 
                    "text": '\n'.join(table_text), 
                    "heading": None
                })

        return elements

    def _split_by_heading(self, elements_list: list) -> list:
        """底层方法：按照 Heading 标题层级物理切割文档"""
        if not elements_list: return []
        result = []
        current_chunk = [elements_list[0]]
        
        for i in range(1, len(elements_list)):
            prev_heading = elements_list[i - 1].get("heading")
            curr_heading = elements_list[i].get("heading")
            
            # 当出现新标题时，切断并开启新的一块
            if curr_heading is not None and prev_heading is None:
                result.append(current_chunk)
                current_chunk = [elements_list[i]]
            else:
                current_chunk.append(elements_list[i])
                
        result.append(current_chunk)
        return result